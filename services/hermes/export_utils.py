"""Export Hermes agent answers to Word, PDF, or PNG files.

Uploads the result to OneDrive and returns the item dict (including webUrl).

Dependencies (add to hermes Dockerfile):
    python-docx
    reportlab
    Pillow

Usage:
    from services.hermes.export_utils import export_answer
    item = export_answer(title, text, fmt, onedrive_client, folder="/Hermes Exports")
    web_url = item.get("webUrl", "")
"""
from __future__ import annotations

import io
import logging
import textwrap
from datetime import datetime
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from services.hermes.onedrive_client import OneDriveClient

logger = logging.getLogger(__name__)

_MAX_LINE_WIDTH = 95  # characters for word-wrap in text formats


def _wrap(text: str) -> str:
    """Wrap long lines for readability in monospace contexts."""
    lines = []
    for para in text.split("\n"):
        if len(para) > _MAX_LINE_WIDTH:
            lines.extend(textwrap.wrap(para, width=_MAX_LINE_WIDTH) or [""])
        else:
            lines.append(para)
    return "\n".join(lines)


# ── Format renderers ──────────────────────────────────────────────────────────

def to_docx(title: str, text: str) -> bytes:
    """Render title + text to a .docx file. Returns bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Date subtitle
    dt_para = doc.add_paragraph(datetime.now().strftime("%Y-%m-%d %H:%M"))
    dt_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    dt_para.runs[0].font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # Body — split on newlines, treating blank lines as paragraph breaks
    current_para_lines: list[str] = []

    def _flush():
        if current_para_lines:
            para = doc.add_paragraph(" ".join(current_para_lines))
            para.runs[0].font.size = Pt(11) if para.runs else None
            current_para_lines.clear()

    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            _flush()
        elif stripped.startswith("# "):
            _flush()
            doc.add_heading(stripped[2:], level=2)
        elif stripped.startswith("## "):
            _flush()
            doc.add_heading(stripped[3:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("• "):
            _flush()
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            current_para_lines.append(stripped)
    _flush()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _register_cjk_font() -> str:
    """Register STSong-Light CIDFont for Chinese support. Returns font name."""
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        return "STSong-Light"
    except Exception:
        return "Helvetica"


def to_pdf(title: str, text: str) -> bytes:
    """Render title + text to a PDF with CJK (Chinese) support. Returns bytes."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    font = _register_cjk_font()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    title_style  = ParagraphStyle("TitleStyle",  parent=styles["Title"],   fontName=font, fontSize=18, spaceAfter=6)
    date_style   = ParagraphStyle("DateStyle",   parent=styles["Normal"],  fontName=font, fontSize=9,
                                   textColor=colors.grey, spaceAfter=12)
    body_style   = ParagraphStyle("BodyStyle",   parent=styles["Normal"],  fontName=font, fontSize=10,
                                   leading=15, spaceAfter=6)
    h2_style     = ParagraphStyle("H2Style",     parent=styles["Heading2"], fontName=font, fontSize=14, spaceAfter=4)
    h3_style     = ParagraphStyle("H3Style",     parent=styles["Heading3"], fontName=font, fontSize=12, spaceAfter=4)
    bullet_style = ParagraphStyle("BulletStyle", parent=styles["Normal"],  fontName=font, fontSize=10,
                                   leading=14, leftIndent=16, spaceAfter=3)

    story = [
        Paragraph(title, title_style),
        Paragraph(datetime.now().strftime("%Y-%m-%d %H:%M"), date_style),
        HRFlowable(width="100%", thickness=0.5, color=colors.grey, spaceAfter=12),
    ]

    current_lines: list[str] = []

    def _flush_para():
        if current_lines:
            raw = " ".join(current_lines)
            story.append(Paragraph(raw, body_style))
            current_lines.clear()

    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            _flush_para()
            story.append(Spacer(1, 4))
        elif stripped.startswith("# "):
            _flush_para()
            story.append(Paragraph(stripped[2:], h2_style))
        elif stripped.startswith("## "):
            _flush_para()
            story.append(Paragraph(stripped[3:], h3_style))
        elif stripped.startswith("- ") or stripped.startswith("• "):
            _flush_para()
            story.append(Paragraph("• " + stripped[2:], bullet_style))
        else:
            current_lines.append(stripped)
    _flush_para()

    doc.build(story)
    return buf.getvalue()


def send_report_as_feishu_pdf(
    title: str,
    text: str,
    open_id: str,
    feishu,
) -> None:
    """Render report to PDF and send it as a Feishu file message.

    Args:
        title:   Report title (used as PDF filename).
        text:    Markdown-ish report body.
        open_id: Feishu recipient open_id.
        feishu:  FeishuClient instance.
    """
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    safe = "".join(c if c.isalnum() or c in " _-" else "_" for c in title)[:60].strip()
    filename = f"{safe}_{ts}.pdf"

    pdf_bytes = to_pdf(title, text)
    file_key  = feishu.upload_file(pdf_bytes, filename, file_type="pdf")
    feishu.send_file(open_id=open_id, file_key=file_key)
    logger.info("Feishu PDF sent: %s (%d bytes)", filename, len(pdf_bytes))


def to_png(title: str, text: str, width: int = 900) -> bytes:
    """Render title + text to a PNG image. Returns bytes."""
    from PIL import Image, ImageDraw, ImageFont

    padding = 40
    line_height = 22
    title_size = 20
    body_size = 14

    # Try to load a decent font; fall back to default
    try:
        font_title = ImageFont.truetype("arial.ttf", title_size)
        font_body  = ImageFont.truetype("arial.ttf", body_size)
        font_date  = ImageFont.truetype("arial.ttf", 11)
    except IOError:
        font_title = ImageFont.load_default()
        font_body  = font_title
        font_date  = font_title

    # Wrap text to fit width
    char_width = body_size * 0.6 or 8
    max_chars = max(int((width - 2 * padding) / char_width), 40)
    wrapped_lines: list[tuple[str, str]] = []  # (line, style)
    wrapped_lines.append((title, "title"))
    wrapped_lines.append((datetime.now().strftime("%Y-%m-%d %H:%M"), "date"))
    wrapped_lines.append(("", "spacer"))

    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            wrapped_lines.append(("", "spacer"))
        else:
            for sub in textwrap.wrap(stripped, width=max_chars) or [""]:
                style = "body"
                if stripped.startswith("# "):
                    sub = sub[2:] if sub.startswith("# ") else sub
                    style = "h2"
                elif stripped.startswith("## "):
                    sub = sub[3:] if sub.startswith("## ") else sub
                    style = "h3"
                wrapped_lines.append((sub, style))

    height = padding * 2 + len(wrapped_lines) * line_height + 10
    img = Image.new("RGB", (width, height), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)

    y = padding
    for line_text, style in wrapped_lines:
        if style == "spacer":
            y += line_height // 2
            continue
        if style == "title":
            draw.text((padding, y), line_text, font=font_title, fill=(30, 30, 30))
            y += title_size + 8
        elif style == "date":
            draw.text((padding, y), line_text, font=font_date, fill=(120, 120, 120))
            y += 18
            draw.line([(padding, y), (width - padding, y)], fill=(200, 200, 200), width=1)
            y += 8
        elif style in ("h2", "h3"):
            draw.text((padding, y), line_text, font=font_title, fill=(50, 80, 130))
            y += title_size + 4
        else:
            draw.text((padding, y), line_text, font=font_body, fill=(40, 40, 40))
            y += line_height

    # Crop to actual content + padding
    final_height = min(y + padding, height)
    img = img.crop((0, 0, width, final_height))

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


# ── Unified export ────────────────────────────────────────────────────────────

def export_answer(
    title: str,
    text: str,
    fmt: str,
    onedrive: "OneDriveClient",
    folder: str = "/Hermes Exports",
) -> dict:
    """Render text to the requested format and upload to OneDrive.

    Args:
        title:    Document title (also used as filename base).
        text:     The answer text (markdown-ish).
        fmt:      "docx" | "pdf" | "png"
        onedrive: Authenticated OneDriveClient instance.
        folder:   OneDrive folder path.

    Returns:
        The OneDrive item dict (contains 'name', 'webUrl', 'id', etc.)
    """
    fmt = fmt.lower().strip()
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    safe_title = "".join(c if c.isalnum() or c in " _-" else "_" for c in title)[:60].strip()
    filename = f"{safe_title}_{ts}.{fmt}"

    if fmt == "docx":
        content_bytes = to_docx(title, text)
    elif fmt == "pdf":
        content_bytes = to_pdf(title, text)
    elif fmt in ("png", "jpg", "jpeg"):
        content_bytes = to_png(title, text)
        filename = f"{safe_title}_{ts}.png"
    else:
        raise ValueError(f"Unsupported format: {fmt}. Use docx, pdf, or png.")

    result = onedrive.upload_file(
        folder_path=folder,
        filename=filename,
        content=content_bytes,
    )
    logger.info("Exported '%s' → OneDrive %s (%d bytes)", filename, result.get("id"), len(content_bytes))
    return result
